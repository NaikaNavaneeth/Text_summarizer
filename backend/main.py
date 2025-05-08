from fastapi import FastAPI, Request, File, UploadFile, Form
from pydantic import BaseModel
from summarizer import summarize_text
from fastapi.middleware.cors import CORSMiddleware
from qa_engine import answer_question
import uvicorn
import logging
from dotenv import load_dotenv
import os
import PyPDF2
import io
import docx

# Load environment variables from .env file
load_dotenv()

# Debug print to check if environment variables are loaded
print(f"GROQ_API_KEY loaded: {'Yes' if os.getenv('GROQ_API_KEY') else 'No'}")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Allow CORS for local frontend testing
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins during development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TextSummaryRequest(BaseModel):
    text: str
    method: str  # "abstractive" or "extractive"
    length: str = "medium"  # "short", "medium", or "detailed"
    useOpenAI: bool = True  # Default to OpenAI if not specified

@app.post("/summarize-text")
async def summarize_text_route(request: TextSummaryRequest):
    logger.info(f"Received summarization request with method: {request.method}, length: {request.length}")
    logger.info(f"Text length: {len(request.text)} characters")
    logger.info(f"Using Groq API")
    
    if len(request.text.strip()) < 10:
        logger.warning("Text too short for meaningful summarization")
        return {"summary": "Text too short for meaningful summarization"}
        
    try:
        summary = summarize_text(request.text, request.method, request.length)
        logger.info(f"Summarization successful, result length: {len(summary)} characters")
        return {"summary": summary}
    except Exception as e:
        logger.error(f"Error in summarization: {str(e)}")
        return {"summary": f"Error: {str(e)}", "error": True}

@app.post("/upload-pdf")
async def upload_pdf(
    file: UploadFile = File(...), 
    method: str = Form(...), 
    length: str = Form("medium"),
    useOpenAI: bool = Form(True)
):
    logger.info(f"Received PDF upload: {file.filename}")
    logger.info(f"Method: {method}, Length: {length}")
    
    try:
        # Read the uploaded file
        contents = await file.read()
        
        # Extract text from PDF
        pdf_text = extract_text_from_pdf(contents)
        
        if len(pdf_text.strip()) < 10:
            logger.warning("Extracted PDF text too short for meaningful summarization")
            return {"summary": "Could not extract sufficient text from the PDF file."}
        
        # Summarize the extracted text
        logger.info(f"Extracted {len(pdf_text)} characters from PDF")
        summary = summarize_text(pdf_text, method, length)
        
        logger.info(f"PDF summarization successful, result length: {len(summary)} characters")
        return {"summary": summary}
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        return {"summary": f"Error processing PDF: {str(e)}", "error": True}

def extract_text_from_pdf(pdf_bytes):
    """Extract text from a PDF file."""
    text = ""
    try:
        # Create a PDF file reader object
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        
        # Get the number of pages in the PDF file
        num_pages = len(pdf_reader.pages)
        
        # Extract text from each page
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

class QARequest(BaseModel):
    context: str
    question: str

@app.post("/ask-question")
async def ask_question_route(request: QARequest):
    logger.info(f"Received question: {request.question}")
    try:
        answer = answer_question(request.context, request.question)
        logger.info("Question answered successfully")
        return {"answer": answer}
    except Exception as e:
        logger.error(f"Error in QA: {str(e)}")
        raise

@app.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...), 
    method: str = Form(...), 
    length: str = Form("medium"),
    useOpenAI: bool = Form(True)
):
    logger.info(f"Received document upload: {file.filename}")
    logger.info(f"Method: {method}, Length: {length}")
    
    try:
        # Read the uploaded file
        contents = await file.read()
        
        # Extract text based on file type
        if file.filename.endswith('.txt'):
            doc_text = contents.decode('utf-8')
        elif file.filename.endswith('.docx'):
            doc_text = extract_text_from_docx(contents)
        else:
            logger.warning(f"Unsupported document type: {file.filename}")
            return {"summary": "Unsupported document type. Please upload a .txt or .docx file."}
        
        if len(doc_text.strip()) < 10:
            logger.warning("Document text too short for meaningful summarization")
            return {"summary": "Document text too short for meaningful summarization"}
        
        # Summarize the extracted text
        logger.info(f"Extracted {len(doc_text)} characters from document")
        summary = summarize_text(doc_text, method, length)
        
        logger.info(f"Document summarization successful, result length: {len(summary)} characters")
        return {"summary": summary}
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return {"summary": f"Error processing document: {str(e)}", "error": True}

def extract_text_from_docx(docx_bytes):
    """Extract text from a DOCX file."""
    text = ""
    try:
        # Create a document object
        doc = docx.Document(io.BytesIO(docx_bytes))
        
        # Extract text from each paragraph
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

# if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8085, reload=True)
